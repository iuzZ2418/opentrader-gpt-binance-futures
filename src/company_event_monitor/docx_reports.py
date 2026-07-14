from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def company_report(snapshot: dict[str, Any]) -> bytes:
    summary = snapshot.get("summary") or snapshot
    company = summary.get("company") or {}
    document = _new_document()
    title = document.add_heading(
        f"{company.get('name', '')}（{company.get('ticker', '')}）公开文本事件报告",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        f"数据截止：{summary.get('data_as_of') or '暂无'}　"
        f"分析窗口：最近 {summary.get('window_days', 180)} 日"
    )
    counts = summary.get("counts") or {}
    document.add_heading("摘要", level=1)
    research_summary = summary.get("research_summary") or {}
    if research_summary.get("summary"):
        document.add_paragraph(research_summary["summary"])
    document.add_paragraph(
        f"发现材料 {summary.get('discovered_document_count', 0)} 条，已分析文档 "
        f"{summary.get('processed_document_count', summary.get('document_count', 0))} 份，"
        "窗口内事件 "
        f"{summary.get('event_count', 0)} 项；正向 {counts.get('positive', 0)} 项，"
        f"负向 {counts.get('negative', 0)} 项，风险升级 {counts.get('escalation', 0)} 项，"
        f"表述冲突 {counts.get('conflict', 0)} 项。"
    )
    if summary.get("source_counts"):
        document.add_paragraph(
            "来源覆盖："
            + "；".join(f"{name} {count} 条" for name, count in summary["source_counts"].items())
        )
    if research_summary.get("research_questions"):
        document.add_heading("后续核验问题", level=1)
        for question in research_summary["research_questions"]:
            document.add_paragraph(question, style="List Bullet")
    workspace = summary.get("research_workspace") or {}
    theses = workspace.get("theses") or []
    if theses:
        document.add_heading("研究观点与证据状态", level=1)
        state_labels = {
            "tracking": "跟踪中",
            "strengthened": "证据加强",
            "weakened": "证据削弱",
            "contested": "证据分歧",
            "confirmed": "人工确认",
            "invalidated": "人工失效",
        }
        for thesis in theses:
            document.add_heading(thesis.get("title") or "未命名观点", level=2)
            state = state_labels.get(thesis.get("state"), thesis.get("state", "跟踪中"))
            document.add_paragraph(
                f"状态：{state}；证据分：{float(thesis.get('evidence_score') or 0):+.2f}；"
                f"支持 {thesis.get('support_count', 0)} 项，反对 "
                f"{thesis.get('contradict_count', 0)} 项。"
            )
            if thesis.get("description"):
                document.add_paragraph(f"观点说明：{thesis['description']}")
            if thesis.get("invalidation_criteria"):
                document.add_paragraph(f"失效条件：{thesis['invalidation_criteria']}")
            for evidence in (thesis.get("evidence") or [])[:5]:
                stance = {
                    "supports": "支持",
                    "contradicts": "反对",
                    "neutral": "中性",
                }.get(evidence.get("stance"), "中性")
                document.add_paragraph(
                    f"[{stance}] {evidence.get('standardized_text', '')} "
                    f"（{str(evidence.get('published_at', ''))[:10]}，"
                    f"{evidence.get('source_name', '')}）",
                    style="List Bullet",
                )
    commitments = workspace.get("commitments") or []
    if commitments:
        document.add_heading("管理层承诺账本", level=1)
        status_labels = {"open": "待验证", "fulfilled": "已兑现", "at_risk": "兑现存疑"}
        for commitment in commitments[:12]:
            document.add_paragraph(
                f"[{status_labels.get(commitment.get('status'), commitment.get('status', ''))}] "
                f"{commitment.get('commitment_text', '')} "
                f"（{str(commitment.get('published_at', ''))[:10]}）",
                style="List Bullet",
            )
    changes = workspace.get("document_changes") or []
    if changes:
        document.add_heading("文档表述变化", level=1)
        for change in changes[:8]:
            document.add_paragraph(
                f"{change.get('document_family', '公告')}："
                f"{change.get('previous_title', '')} → {change.get('current_title', '')}；"
                f"文本相似度 {float(change.get('similarity') or 0):.1%}。",
                style="List Bullet",
            )
            if change.get("added_text"):
                document.add_paragraph(f"新增表述：{str(change['added_text'])[:600]}")
    market = summary.get("market_analysis") or {}
    if market:
        document.add_heading("价格与事件联动", level=1)
        returns = market.get("returns") or {}
        benchmark = market.get("benchmark") or {}
        risk = market.get("risk") or {}
        document.add_paragraph(
            f"行情截止：{market.get('as_of', '暂无')}　最新价格："
            f"{float(market.get('latest_price') or 0):.2f}　"
            f"近20日收益：{_percent(returns.get('20d'))}　"
            f"近60日收益：{_percent(returns.get('60d'))}　"
            f"相对{benchmark.get('name', '基准')}20日超额："
            f"{_percent(benchmark.get('excess_20d'))}"
        )
        document.add_paragraph(
            f"20日年化波动：{_probability(risk.get('annualized_volatility_20d'))}　"
            f"60日最大回撤：{_percent(risk.get('max_drawdown_60d'))}"
        )
        valuation = market.get("valuation") or {}
        if valuation:
            document.add_paragraph(
                f"估值快照：动态PE {_multiple(valuation.get('pe_dynamic'))}　"
                f"PB {_multiple(valuation.get('pb'))}　"
                f"总市值 {_number(valuation.get('total_market_cap_yi'))} 亿元　"
                f"换手率 {_source_percent(valuation.get('turnover_rate'))}"
            )
        forecast = market.get("forecast_20d") or {}
        probabilities = forecast.get("probabilities") or {}
        price_range = forecast.get("price_range") or {}
        document.add_paragraph(
            f"20交易日研究情景：{forecast.get('regime', '暂无')}；"
            f"上行/震荡/下行概率分别为 {_probability(probabilities.get('up'))} / "
            f"{_probability(probabilities.get('neutral'))} / "
            f"{_probability(probabilities.get('down'))}。"
        )
        document.add_paragraph(
            f"历史分布价格区间：下行情景 {float(price_range.get('downside_p10') or 0):.2f}，"
            f"中位情景 {float(price_range.get('median_p50') or 0):.2f}，"
            f"上行情景 {float(price_range.get('upside_p90') or 0):.2f}。"
        )
        document.add_paragraph(market.get("disclaimer") or "价格情景不构成投资建议。")
        links = market.get("event_price_links") or []
        if links:
            document.add_heading("公告后价格反应", level=2)
            reaction_table = document.add_table(rows=1, cols=5)
            reaction_table.style = "Table Grid"
            for cell, label in zip(
                reaction_table.rows[0].cells,
                ("公告日", "标准化事件", "后5日", "后20日", "20日超额"),
                strict=True,
            ):
                cell.text = label
            for link in links[:12]:
                values = (
                    link.get("date", ""),
                    link.get("event", ""),
                    _percent(link.get("forward_5d")),
                    _percent(link.get("forward_20d")),
                    _percent(link.get("excess_20d")),
                )
                for cell, value in zip(
                    reaction_table.add_row().cells,
                    values,
                    strict=True,
                ):
                    cell.text = str(value)
    document.add_heading("重要事件与证据", level=1)
    for event in summary.get("important_events") or []:
        document.add_heading(event.get("standardized_text") or event.get("event_type", "事件"), 2)
        document.add_paragraph(
            f"变化：{event.get('change_type', '')}　研究关注度："
            f"{float(event.get('value_score') or 0):.2f}　来源：{event.get('source_name', '')}"
        )
        if event.get("importance_reason"):
            document.add_paragraph(f"重要性：{event['importance_reason']}")
        document.add_paragraph(f"原文证据：{event.get('evidence_text', '')}")
        location = []
        if event.get("evidence_page"):
            location.append(f"第 {event['evidence_page']} 页")
        if event.get("evidence_section"):
            location.append(str(event["evidence_section"]))
        document.add_paragraph(
            f"发布日期：{event.get('published_at', '')}　位置：{' / '.join(location) or '正文'}　"
            f"链接：{event.get('url', '')}"
        )
    document.add_paragraph("本报告仅用于公开信息研究辅助，不构成投资建议或自动评级。")
    return _save(document)


def comparison_report(comparison: dict[str, Any]) -> bytes:
    result = comparison.get("result") or comparison
    document = _new_document()
    title = document.add_heading("上市公司公开文本横向比较报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if result.get("partial_update_failure"):
        document.add_paragraph("注意：部分公司更新失败，报告已明确列示各公司实际数据截止日期。")
    document.add_heading("比较摘要", level=1)
    document.add_paragraph(result.get("deepseek_summary") or "暂无摘要")
    document.add_heading("公司横向总表", level=1)
    companies = result.get("companies") or []
    table = document.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    for cell, label in zip(
        table.rows[0].cells,
        (
            "公司",
            "代码",
            "数据截止",
            "文档",
            "事件",
            "正向",
            "负向",
            "20日收益",
            "相对基准",
            "20日情景",
        ),
        strict=True,
    ):
        cell.text = label
    for company in companies:
        cells = table.add_row().cells
        market = company.get("price_analysis") or {}
        forecast = market.get("forecast_20d") or {}
        values = (
            company.get("name", ""),
            company.get("ticker", ""),
            company.get("data_as_of", ""),
            company.get("document_count", 0),
            company.get("event_count", 0),
            company.get("positive", 0),
            company.get("negative", 0),
            _percent((market.get("returns") or {}).get("20d")),
            _percent((market.get("benchmark") or {}).get("excess_20d")),
            forecast.get("regime", "暂无"),
        )
        for cell, value in zip(cells, values, strict=True):
            cell.text = str(value)
    document.add_heading("各公司最重要的三项变化", level=1)
    for company in companies:
        document.add_heading(f"{company.get('name')}（{company.get('ticker')}）", level=2)
        for event in company.get("top_events") or []:
            document.add_paragraph(
                f"{event.get('standardized_text', '')}\n"
                f"原文：{event.get('evidence_text', '')}\n"
                f"来源：{event.get('source_name', '')} {event.get('published_at', '')} "
                f"{event.get('url', '')}"
            )
    document.add_heading("共同风险与差异化事件", level=1)
    document.add_paragraph("共同风险：" + "、".join(result.get("common_risks") or ["暂无"]))
    document.add_paragraph("差异化事件：" + "、".join(result.get("distinctive_events") or ["暂无"]))
    document.add_paragraph(result.get("disclaimer") or "本报告不构成投资建议。")
    return _save(document)


def _new_document() -> Document:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    return document


def _save(document: Document) -> bytes:
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def _percent(value: Any) -> str:
    try:
        return f"{float(value):+.1%}"
    except (TypeError, ValueError):
        return "暂无"


def _source_percent(value: Any) -> str:
    try:
        return f"{float(value):.2f}%"
    except (TypeError, ValueError):
        return "暂无"


def _probability(value: Any) -> str:
    try:
        return f"{float(value):.1%}"
    except (TypeError, ValueError):
        return "暂无"


def _multiple(value: Any) -> str:
    try:
        return f"{float(value):.2f}x"
    except (TypeError, ValueError):
        return "暂无"


def _number(value: Any) -> str:
    try:
        return f"{float(value):,.2f}"
    except (TypeError, ValueError):
        return "暂无"
